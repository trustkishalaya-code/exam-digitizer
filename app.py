import io
import os
import streamlit as st
from PIL import Image, ImageEnhance, ImageOps
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import google.generativeai as genai

# --- PAGE CONFIGURATION ---
st.set_page_config(
    page_title="ACADEMIC STUDIO // ELECTRIC YELLOW BRUTALISM",
    page_icon="⚡",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# --- ELECTRIC YELLOW BRUTALIST CSS THEME ---
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Space+Grotesk:wght@400;600;700&family=JetBrains+Mono:wght@500;800&display=swap');

    :root {
        --bg-color: #f4f4f0;
        --electric-yellow: #ccff00;
        --pure-black: #0d0d0d;
        --pure-white: #ffffff;
        --border-width: 3px;
        --shadow-offset: 4px;
    }

    .stApp {
        background-color: var(--bg-color);
        font-family: 'Space Grotesk', sans-serif;
        color: var(--pure-black);
    }

    #MainMenu, header, footer {visibility: hidden;}

    .brutalist-header {
        background-color: var(--electric-yellow);
        border: var(--border-width) solid var(--pure-black);
        padding: 2rem;
        box-shadow: var(--shadow-offset) var(--shadow-offset) 0px var(--pure-black);
        margin-bottom: 2rem;
        display: flex;
        justify-content: space-between;
        align-items: center;
    }
    
    .brutalist-header h1 {
        font-family: 'JetBrains Mono', monospace;
        font-weight: 800;
        font-size: 2.5rem;
        color: var(--pure-black);
        margin: 0;
        text-transform: uppercase;
        letter-spacing: -1px;
    }

    .brutalist-header p {
        font-family: 'JetBrains Mono', monospace;
        font-weight: 500;
        margin: 0;
        background: var(--pure-black);
        color: var(--electric-yellow);
        padding: 0.4rem 0.8rem;
        font-size: 0.9rem;
        border: 2px solid var(--pure-black);
    }

    .bento-card {
        background-color: var(--pure-white);
        border: var(--border-width) solid var(--pure-black);
        padding: 1.5rem;
        box-shadow: var(--shadow-offset) var(--shadow-offset) 0px var(--pure-black);
        margin-bottom: 1.5rem;
        position: relative;
    }

    .bento-card-yellow {
        background-color: var(--electric-yellow);
        border: var(--border-width) solid var(--pure-black);
        padding: 1.5rem;
        box-shadow: var(--shadow-offset) var(--shadow-offset) 0px var(--pure-black);
        margin-bottom: 1.5rem;
    }

    .card-title {
        font-family: 'JetBrains Mono', monospace;
        font-weight: 800;
        font-size: 1.2rem;
        text-transform: uppercase;
        border-bottom: var(--border-width) solid var(--pure-black);
        padding-bottom: 0.5rem;
        margin-bottom: 1rem;
        display: flex;
        align-items: center;
        gap: 0.5rem;
    }

    .stButton > button {
        background-color: var(--electric-yellow) !important;
        color: var(--pure-black) !important;
        font-family: 'JetBrains Mono', monospace !important;
        font-weight: 800 !important;
        border: var(--border-width) solid var(--pure-black) !important;
        padding: 0.75rem 1.5rem !important;
        box-shadow: var(--shadow-offset) var(--shadow-offset) 0px var(--pure-black) !important;
        border-radius: 0px !important;
        width: 100% !important;
        text-transform: uppercase !important;
        transition: transform 0.1s ease, box-shadow 0.1s ease !important;
    }

    .stButton > button:hover {
        transform: translate(-2px, -2px);
        box-shadow: calc(var(--shadow-offset) + 2px) calc(var(--shadow-offset) + 2px) 0px var(--pure-black) !important;
    }

    .stButton > button:active {
        transform: translate(2px, 2px);
        box-shadow: 2px 2px 0px var(--pure-black) !important;
    }

    input, textarea, select {
        border: var(--border-width) solid var(--pure-black) !important;
        border-radius: 0px !important;
        background-color: var(--pure-white) !important;
        font-family: 'Space Grotesk', sans-serif !important;
        font-weight: 600 !important;
        box-shadow: 3px 3px 0px rgba(13, 13, 13, 0.2) !important;
    }

    div[data-testid="metric-container"] {
        background-color: var(--pure-white);
        border: var(--border-width) solid var(--pure-black);
        padding: 1rem;
        box-shadow: var(--shadow-offset) var(--shadow-offset) 0px var(--pure-black);
    }
    
    div[data-testid="metric-container"] label {
        font-family: 'JetBrains Mono', monospace !important;
        font-weight: 800 !important;
        text-transform: uppercase;
    }
</style>
""", unsafe_allow_html=True)

# --- GEMINI AI & OCR LOGIC ---
def extract_questions_with_gemini(image, api_key):
    if not api_key:
        return None, "API Key missing."
    try:
        genai.configure(api_key=api_key)
        # Using fallback model logic robustly
        model_name = 'gemini-1.5-pro'
        try:
            model = genai.GenerativeModel(model_name)
        except Exception:
            model = genai.GenerativeModel('gemini-1.5-flash')

        prompt = """
        Analyze this handwritten or printed exam paper image. Extract all questions, their types (Descriptive, MCQ, Fill-in-the-blanks, etc.), allocated marks, and any options (if MCQ). 
        Return the result strictly as a structured list or JSON-like breakdown if possible, or clear text format containing:
        - Question Number / Text
        - Type
        - Marks
        - Options (if applicable)
        """
        response = model.generate_content([prompt, image])
        return response.text, None
    except Exception as e:
        return None, str(e)

# --- WORD DOCUMENT GENERATOR CORE ---
def create_word_document(metadata, questions, layout_blocks):
    doc = Document()
    
    # Page setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        if metadata.get("orientation") == "Landscape":
            section.orientation = WD_ALIGN_PARAGRAPH.CENTER # Handled via section props if needed
            section.page_width = Inches(11.0)
            section.page_height = Inches(8.5)

    # Header / Metadata Section
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(metadata.get("institution", "ACADEMIC INSTITUTION").upper())
    title_run.font.name = metadata.get("font_family", "Kalpurush")
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 0)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run(f"{metadata.get('exam_name', 'EXAMINATION')} — {metadata.get('academic_year', '2026')}")
    sub_run.font.name = metadata.get("font_family", "Kalpurush")
    sub_run.font.size = Pt(13)
    sub_run.font.bold = True

    # Exam Info Meta Table
    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data_labels = [
        [f"Subject: {metadata.get('subject', 'General')}", f"Full Marks: {metadata.get('full_marks', '100')}"],
        [f"Class: {metadata.get('class_name', 'Primary')}", f"Duration: {metadata.get('duration', '2 Hours')}"]
    ]
    for r_idx, row in enumerate(meta_table.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.text = meta_data_labels[r_idx][c_idx]
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = metadata.get("font_family", "Kalpurush")
                    run.font.size = Pt(11)
                    run.font.bold = True

    doc.add_paragraph() 

    # Layout Blocks Insertion
    if layout_blocks:
        bp = doc.add_paragraph()
        b_run = bp.add_run(f"Special Instructions / Layout Focus: {', '.join(layout_blocks)}")
        b_run.font.name = metadata.get("font_family", "Kalpurush")
        b_run.font.size = Pt(10)
        b_run.font.italic = True
        doc.add_paragraph()

    # Questions Loop
    for idx, q in enumerate(questions, 1):
        qp = doc.add_paragraph()
        q_run = qp.add_run(f"Q{idx}. {q.get('text', '')} ")
        q_run.font.name = metadata.get("font_family", "Kalpurush")
        q_run.font.size = Pt(11)
        q_run.font.bold = True
        
        marks_run = qp.add_run(f"[{q.get('marks', '5')} Marks]")
        marks_run.font.name = metadata.get("font_family", "Kalpurush")
        marks_run.font.size = Pt(10)
        marks_run.font.italic = True

        if q.get('type') == 'MCQ':
            for opt in ['A', 'B', 'C', 'D']:
                opt_val = q.get(f'opt_{opt.lower()}', '')
                if opt_val:
                    opt_p = doc.add_paragraph()
                    opt_run = opt_p.add_run(f"   ({opt}) {opt_val}")
                    opt_run.font.name = metadata.get("font_family", "Kalpurush")
                    opt_run.font.size = Pt(11)

        doc.add_paragraph() 

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

# --- UI LAYOUT ---

st.markdown("""
    <div class="brutalist-header">
        <div>
            <h1>Academic Studio</h1>
            <p>// High-Precision Exam Paper Digitizer & Layout Compiler</p>
        </div>
        <div style="font-family: 'JetBrains Mono', monospace; font-size: 1.5rem; font-weight: 800; border: 3px solid #0d0d0d; background: #ffffff; padding: 0.5rem 1rem;">
            v2.6.0
        </div>
    </div>
""", unsafe_allow_html=True)

# Session state initialization
if 'questions' not in st.session_state:
    st.session_state.questions = [
        {"type": "Descriptive", "text": "Describe the core ecosystem dynamics discussed in class.", "marks": "10"},
        {"type": "MCQ", "text": "Which layout paradigm optimizes structural density?", "marks": "5", "opt_a": "Editorial Minimalist", "opt_b": "Bento Grid", "opt_c": "Electric Brutalism", "opt_d": "All of the above"}
    ]

col1, col2 = st.columns([1, 1], gap="large")

with col1:
    # 1. Structural Metadata Editor
    st.markdown('<div class="bento-card">', unsafe_allow_html=True)
    st.markdown('<div class="card-title">⚙️ Structural Metadata Editor</div>', unsafe_allow_html=True)
    
    institution = st.text_input("Institution Name", "BANGLADESH INTERNATIONAL SCHOOL")
    exam_name = st.text_input("Examination Name", "TERM FINAL EXAMINATION")
    
    c_meta1, c_meta2 = st.columns(2)
    with c_meta1:
        subject = st.text_input("Subject", "Mathematics & Logic")
        class_name = st.text_input("Class Level", "Grade 5")
        orientation = st.selectbox("Page Orientation", ["Portrait", "Landscape"])
    with c_meta2:
        academic_year = st.text_input("Academic Year", "2026")
        full_marks = st.text_input("Full Marks", "100")
        duration = st.text_input("Duration", "2.5 Hours")

    font_family = st.selectbox("Typography Preset", ["Kalpurush", "Mangal", "Times New Roman", "Arial"])
    layout_blocks = st.multiselect("Active Layout Blocks", ["Drawing Box Area", "Two-Column Split", "Answer Lines Space", "Watermark Notice"], default=["Answer Lines Space"])
    st.markdown('</div>', unsafe_allow_html=True)

    # 2. Handwriting Enhancement & Vision AI OCR Studio
    st.markdown('<div class="bento-card-yellow">', unsafe_allow_html=True)
    st.markdown('<div class="card-title">👁️ Handwriting OCR & Vision AI Studio</div>', unsafe_allow_html=True)
    
    gemini_api_key = st.text_input("Gemini API Key", type="password", value=os.environ.get("GEMINI_API_KEY", ""))
    uploaded_image = st.file_uploader("Upload Exam Paper Scan / Image", type=["png", "jpg", "jpeg"])

    if uploaded_image:
        image = Image.open(uploaded_image)
        st.image(image, caption="Uploaded Scan Preview", use_container_width=True)
        
        c_enh1, c_enh2 = st.columns(2)
        with c_enh1:
            contrast_val = st.slider("Contrast Enhancement", 0.5, 3.0, 1.5)
        with c_enh2:
            sharpness_val = st.slider("Sharpness Enhancement", 0.5, 3.0, 2.0)
            
        # Apply preprocessing
        enhancer = ImageEnhance.Contrast(image)
        img_processed = enhancer.enhance(contrast_val)
        sharp_enhancer = ImageEnhance.Sharpness(img_processed)
        img_processed = sharp_enhancer.enhance(sharpness_val)
        
        st.image(img_processed, caption="Preprocessed Image (OCR Ready)", use_container_width=True)

        if st.button("🚀 Extract Questions via Gemini OCR"):
            with st.spinner("Analyzing handwriting patterns and compiling structures..."):
                ocr_result, err = extract_questions_with_gemini(img_processed, gemini_api_key)
                if err:
                    st.error(f"OCR Error: {err}")
                else:
                    st.success("Extraction Complete!")
                    st.text_area("Raw AI Extraction Output", ocr_result, height=150)
    st.markdown('</div>', unsafe_allow_html=True)

    # 3. Question Architecture Form
    st.markdown('<div class="bento-card">', unsafe_allow_html=True)
    st.markdown('<div class="card-title">📝 Manual Question Architecture</div>', unsafe_allow_html=True)
    
    q_type = st.selectbox("Question Type", ["Descriptive", "MCQ", "Fill-in-the-Blanks"])
    q_text = st.text_area("Question Prompt", placeholder="Enter exam question statement...")
    q_marks = st.text_input("Allocated Marks", "5")
    
    opt_a, opt_b, opt_c, opt_d = "", "", "", ""
    if q_type == "MCQ":
        oc1, oc2 = st.columns(2)
        with oc1:
            opt_a = st.text_input("Option A", "")
            opt_b = st.text_input("Option B", "")
        with oc2:
            opt_c = st.text_input("Option C", "")
            opt_d = st.text_input("Option D", "")

    if st.button("➕ Inject Question to Stack"):
        if q_text:
            new_q = {"type": q_type, "text": q_text, "marks": q_marks}
            if q_type == "MCQ":
                new_q.update({"opt_a": opt_a, "opt_b": opt_b, "opt_c": opt_c, "opt_d": opt_d})
            st.session_state.questions.append(new_q)
            st.rerun()
        else:
            st.error("Question prompt cannot be empty.")
    st.markdown('</div>', unsafe_allow_html=True)

with col2:
    # 4. Active Document Pipeline & Queue
    st.markdown('<div class="bento-card">', unsafe_allow_html=True)
    st.markdown('<div class="card-title">📊 Active Document Pipeline</div>', unsafe_allow_html=True)
    
    st.metric("Total Compiled Questions", len(st.session_state.questions))
    
    st.markdown("### Current Queue Preview")
    for i, q in enumerate(st.session_state.questions, 1):
        st.markdown(f"""
        <div style="border: 2px solid #0d0d0d; padding: 0.75rem; margin-bottom: 0.5rem; background: #f9f9f6;">
            <strong>Q{i} [{q['type']}]:</strong> {q['text']} <br>
            <span style="font-size: 0.85rem; font-family: 'JetBrains Mono'; background: #ccff00; padding: 2px 6px; border: 1px solid #0d0d0d;">Marks: {q['marks']}</span>
        </div>
        """, unsafe_allow_html=True)
        
    if st.button("🗑️ Clear Question Queue"):
        st.session_state.questions = []
        st.rerun()
        
    st.markdown('</div>', unsafe_allow_html=True)

    # 5. Compile & Export Studio
    st.markdown('<div class="bento-card-yellow">', unsafe_allow_html=True)
    st.markdown('<div class="card-title">💾 Compile & Export Word Document</div>', unsafe_allow_html=True)
    st.write("Generate a professionally formatted Word document ready for print distribution with your selected font preset, layout blocks, and structured metadata.")
    
    metadata = {
        "institution": institution,
        "exam_name": exam_name,
        "subject": subject,
        "class_name": class_name,
        "academic_year": academic_year,
        "full_marks": full_marks,
        "duration": duration,
        "font_family": font_family,
        "orientation": orientation
    }

    if st.session_state.questions:
        doc_bytes = create_word_document(metadata, st.session_state.questions, layout_blocks)
        st.download_button(
            label="⚡ DOWNLOAD COMPILED .DOCX",
            data=doc_bytes,
            file_name=f"{subject.replace(' ', '_')}_Exam_Paper.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.warning("Add at least one question to enable document compilation.")
        
    st.markdown('</div>', unsafe_allow_html=True)
